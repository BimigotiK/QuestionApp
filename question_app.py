#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
QuestionApp – приложение для работы с Word‑файлами на PyQt5.
"""

import random
import io
import os
import sys
import tempfile
import logging
import hashlib
import re
from pathlib import Path
from PIL import Image, ImageQt
from functools import lru_cache

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QScrollArea, QCheckBox, QFrame,
    QFileDialog, QMessageBox, QGroupBox, QSizePolicy, QMenu, QAction,
    QLineEdit, QComboBox, QProgressDialog, QTabWidget, QDialog,
    QDialogButtonBox, QShortcut, QSpinBox, QRadioButton, QButtonGroup
)
from PyQt5.QtCore import (
    Qt, QSize, pyqtSignal, QTranslator, QLocale, QThread,
    QSettings
)
from PyQt5.QtGui import (
    QPixmap, QFont, QDragEnterEvent, QDropEvent, QImage,
    QPalette, QColor, QKeySequence
)

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Настройка логирования
logging.basicConfig(
    level=logging.WARNING,  # Изменено с INFO на WARNING
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('question_app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


class ImageCache:
    """Класс для кэширования и управления изображениями."""
    
    def __init__(self):
        self.cache = {}
        self.hash_cache = {}
        
    def get_image_hash(self, img_bytes):
        """Получает хэш изображения для использования в качестве ключа кэша."""
        return hashlib.md5(img_bytes).hexdigest()
    
    @lru_cache(maxsize=100)
    def get_scaled_pixmap(self, img_hash, max_width=600, max_height=400):
        """Возвращает масштабированное изображение из кэша."""
        if img_hash in self.cache:
            return self.cache[img_hash]
        return None
    
    def scale_and_cache_image(self, img_bytes, max_width=600, max_height=400):
        """Масштабирует и кэширует изображение."""
        img_hash = self.get_image_hash(img_bytes)
        
        # Проверяем кэш
        cached = self.get_scaled_pixmap(img_hash, max_width, max_height)
        if cached:
            return cached, img_hash
        
        try:
            # Загружаем изображение через PIL
            image = Image.open(io.BytesIO(img_bytes))
            
            # Конвертируем в RGB если нужно
            if image.mode not in ['RGB', 'RGBA']:
                image = image.convert('RGB')
            
            # Масштабируем изображение с сохранением пропорций
            width, height = image.size
            if width > max_width or height > max_height:
                ratio = min(max_width/width, max_height/height)
                new_size = (int(width * ratio), int(height * ratio))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            # Конвертируем PIL Image в QImage
            if image.mode == "RGB":
                qimage = QImage(image.tobytes(), image.size[0], image.size[1], 
                               image.size[0] * 3, QImage.Format_RGB888)
            elif image.mode == "RGBA":
                qimage = QImage(image.tobytes(), image.size[0], image.size[1],
                               image.size[0] * 4, QImage.Format_RGBA8888)
            
            pixmap = QPixmap.fromImage(qimage)
            
            # Кэшируем
            self.cache[img_hash] = pixmap
            self.hash_cache[img_hash] = img_bytes
            
            return pixmap, img_hash
            
        except Exception as e:
            logger.error(f"Ошибка обработки изображения: {str(e)}")
            return None, None
    
    def clear_cache(self):
        """Очищает кэш изображений."""
        self.cache.clear()
        self.hash_cache.clear()
        self.get_scaled_pixmap.cache_clear()


class ThemeManager:
    """Менеджер тем оформления."""
    
    themes = {
        'light': {
            'name': 'Светлая',
            'primary_bg': '#ffffff',
            'secondary_bg': '#f5f5f5',
            'tertiary_bg': '#f0f0f0',
            'primary_text': '#333333',
            'secondary_text': '#666666',
            'accent': '#4CAF50',
            'accent_hover': '#45a049',
            'border': '#dddddd',
            'hover_bg': '#f0f8ff',
            'selected_bg': '#e8f5e9',
            'button_bg': '#4CAF50',
            'button_hover': '#45a049',
            'scrollbar_bg': '#f0f0f0',
            'scrollbar_handle': '#c1c1c1',
            'error': '#dc3545',
            'warning': '#ffc107',
            'success': '#28a745',
            'info': '#17a2b8',
            'drag_drop_bg': '#f0f0f0',
            'drag_drop_border': '#aaa',
            'drag_drop_text': '#333',
            'checkbox_border': '#333',
            'separator_color': '#ccc'
        },
        'dark': {
            'name': 'Темная',
            'primary_bg': '#2b2b2b',
            'secondary_bg': '#3c3c3c',
            'tertiary_bg': '#4a4a4a',
            'primary_text': '#e0e0e0',
            'secondary_text': '#aaaaaa',
            'accent': '#66bb6a',
            'accent_hover': '#5cb860',
            'border': '#666666',
            'hover_bg': '#3a4a5a',
            'selected_bg': '#2d4a2d',
            'button_bg': '#66bb6a',
            'button_hover': '#5cb860',
            'scrollbar_bg': '#3c3c3c',
            'scrollbar_handle': '#888888',
            'error': '#ff5252',
            'warning': '#ffb74d',
            'success': '#81c784',
            'info': '#4fc3f7',
            'drag_drop_bg': '#3c3c3c',
            'drag_drop_border': '#888888',
            'drag_drop_text': '#e0e0e0',
            'checkbox_border': '#e0e0e0',
            'separator_color': '#666666'
        }
    }
    
    def __init__(self):
        self.current_theme = 'light'
        
    def get_theme(self, theme_name=None):
        """Возвращает словарь с цветами темы."""
        if theme_name is None:
            theme_name = self.current_theme
        return self.themes.get(theme_name, self.themes['light'])
    
    def apply_theme(self, theme_name, app):
        """Применяет тему к приложению."""
        self.current_theme = theme_name
        theme = self.get_theme(theme_name)
        
        # Создаем палитру для приложения
        palette = QPalette()
        
        # Базовые цвета
        palette.setColor(QPalette.Window, QColor(theme['primary_bg']))
        palette.setColor(QPalette.WindowText, QColor(theme['primary_text']))
        palette.setColor(QPalette.Base, QColor(theme['secondary_bg']))
        palette.setColor(QPalette.AlternateBase, QColor(theme['tertiary_bg']))
        palette.setColor(QPalette.ToolTipBase, QColor(theme['primary_bg']))
        palette.setColor(QPalette.ToolTipText, QColor(theme['primary_text']))
        palette.setColor(QPalette.Text, QColor(theme['primary_text']))
        palette.setColor(QPalette.Button, QColor(theme['secondary_bg']))
        palette.setColor(QPalette.ButtonText, QColor(theme['primary_text']))
        palette.setColor(QPalette.BrightText, QColor(theme['accent']))
        palette.setColor(QPalette.Link, QColor(theme['accent']))
        palette.setColor(QPalette.Highlight, QColor(theme['accent']))
        palette.setColor(QPalette.HighlightedText, QColor(theme['primary_bg']))
        
        app.setPalette(palette)
        
        # Стили для виджетов
        style_sheet = f"""
            QMainWindow, QWidget {{
                background-color: {theme['primary_bg']};
                color: {theme['primary_text']};
            }}
            
            QGroupBox {{
                font-weight: bold;
                border: 2px solid {theme['border']};
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
                background-color: {theme['secondary_bg']};
            }}
            
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: {theme['primary_text']};
            }}
            
            QPushButton {{
                background-color: {theme['button_bg']};
                color: white;
                border: none;
                padding: 8px;
                border-radius: 5px;
                font-weight: bold;
            }}
            
            QPushButton:hover {{
                background-color: {theme['button_hover']};
            }}
            
            QPushButton:pressed {{
                background-color: {theme['accent_hover']};
            }}
            
            QCheckBox {{
                spacing: 5px;
                color: {theme['primary_text']};
            }}
            
            QCheckBox::indicator {{
                width: 20px;
                height: 20px;
                border: 2px solid {theme['checkbox_border']};
                border-radius: 3px;
                background: {theme['primary_bg']};
            }}
            
            QCheckBox::indicator:checked {{
                background-color: {theme['accent']};
                border: 2px solid {theme['accent']};
            }}
            
            QScrollArea {{
                border: none;
                background-color: {theme['primary_bg']};
            }}
            
            QScrollBar:vertical {{
                background-color: {theme['scrollbar_bg']};
                width: 12px;
                margin: 0px;
            }}
            
            QScrollBar::handle:vertical {{
                background-color: {theme['scrollbar_handle']};
                min-height: 20px;
                border-radius: 6px;
            }}
            
            QScrollBar::handle:vertical:hover {{
                background-color: {theme['accent']};
            }}
            
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
                height: 0px;
            }}
            
            QLineEdit, QComboBox, QSpinBox {{
                border: 1px solid {theme['border']};
                border-radius: 4px;
                padding: 5px;
                background-color: {theme['primary_bg']};
                color: {theme['primary_text']};
            }}
            
            QLineEdit:focus, QComboBox:focus, QSpinBox:focus {{
                border: 2px solid {theme['accent']};
            }}
            
            QProgressBar {{
                border: 1px solid {theme['border']};
                border-radius: 4px;
                text-align: center;
                background-color: {theme['secondary_bg']};
            }}
            
            QProgressBar::chunk {{
                background-color: {theme['accent']};
                border-radius: 4px;
            }}
            
            QRadioButton {{
                color: {theme['primary_text']};
                spacing: 5px;
            }}
            
            QLabel {{
                color: {theme['primary_text']};
            }}
            
            QFrame[frameShape="4"] {{  /* HLine */
                color: {theme['separator_color']};
                background-color: {theme['separator_color']};
            }}
        """
        
        app.setStyleSheet(style_sheet)
        
        return style_sheet


class ExportWorker(QThread):
    """Поток для экспорта в фоновом режиме."""
    progress = pyqtSignal(int)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)
    
    def __init__(self, questions, export_format, export_path, export_options):
        super().__init__()
        self.questions = questions
        self.export_format = export_format
        self.export_path = export_path
        self.export_options = export_options
        
    def run(self):
        try:
            if self.export_format == 'docx':
                self.export_to_docx()
            elif self.export_format == 'txt':
                self.export_to_txt()
            elif self.export_format == 'html':
                self.export_to_html()
            elif self.export_format == 'json':
                self.export_to_json()
            elif self.export_format == 'pdf':
                self.export_to_pdf()
                
            self.finished.emit(self.export_path)
            
        except Exception as e:
            logger.error(f"Ошибка экспорта: {str(e)}")
            self.error.emit(str(e))
    
    def export_to_docx(self):
        """Экспорт в DOCX с сохранением форматирования."""
        doc = Document()
        
        # Устанавливаем поля документа
        sections = doc.sections
        for section in sections:
            section.left_margin = Cm(3)      # 3 см слева
            section.right_margin = Cm(1)     # 1 см справа
            section.top_margin = Cm(1)       # 1 см сверху
            section.bottom_margin = Cm(1)    # 1 см снизу
        
        # Настраиваем стили
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        
        # Высота страницы A4 в сантиметрах и расчет 25%
        PAGE_HEIGHT_CM = 29.7
        MAX_HEIGHT_CM = PAGE_HEIGHT_CM * 0.25  # 25% от высоты страницы
        
        # Переводим сантиметры в дюймы (1 дюйм = 2.54 см)
        MAX_HEIGHT_INCHES = MAX_HEIGHT_CM / 2.54
        
        for idx, question in enumerate(self.questions):
            # Получаем текст вопроса
            question_text = question['text']
            images = question['images']
            
            # Если выбрана нумерация по порядку, заменяем номер в тексте
            if self.export_options.get('numbering') == 'sequential':
                # Ищем и заменяем "Aufgabe XX" на "Aufgabe [номер]"
                pattern = r'(Aufgabe\s+)(\d+)'
                new_number = str(idx + 1)
                question_text = re.sub(pattern, rf'\g<1>{new_number}', question_text, flags=re.IGNORECASE)
            
            # Разделяем текст на части по маркерам [BILD]
            text_parts = question_text.split('[BILD]')
            
            for i, part in enumerate(text_parts):
                if part.strip():
                    # Сохраняем оригинальное форматирование с пробелами
                    lines = part.split('\n')
                    for line in lines:
                        if line or (not line and i > 0):  # Сохраняем пустые строки
                            p = doc.add_paragraph(line.rstrip())
                            # Сохраняем оригинальные отступы
                            if line.startswith('    ') or line.startswith('\t'):
                                p.paragraph_format.left_indent = Pt(20)
                
                # Добавляем изображение
                if i < len(images):
                    try:
                        img_bytes = images[i]
                        
                        # Загружаем изображение для проверки размеров
                        img = Image.open(io.BytesIO(img_bytes))
                        img_width, img_height = img.size
                        
                        # Переводим пиксели в дюймы (при 96 DPI)
                        # 1 дюйм = 96 пикселей при 96 DPI
                        height_inches = img_height / 96.0
                        
                        # Если высота изображения больше 25% страницы, уменьшаем
                        if height_inches > MAX_HEIGHT_INCHES:
                            # Вычисляем коэффициент масштабирования
                            scale_factor = MAX_HEIGHT_INCHES / height_inches
                            
                            # Вычисляем новые размеры в пикселях
                            new_width = int(img_width * scale_factor)
                            new_height = int(img_height * scale_factor)
                            
                            # Изменяем размер изображения
                            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                            
                            # Сохраняем измененное изображение в буфер
                            img_buffer = io.BytesIO()
                            img.save(img_buffer, format='PNG')
                            img_buffer.seek(0)
                            
                            # Добавляем изображение с ограниченным размером
                            doc.add_picture(img_buffer, width=Inches(MAX_HEIGHT_INCHES * (img_width/img_height)))
                            
                            logger.info(f"Изображение уменьшено: {height_inches:.2f} → {MAX_HEIGHT_INCHES:.2f} дюймов")
                        else:
                            # Добавляем изображение оригинального размера
                            img_buffer = io.BytesIO(img_bytes)
                            img_buffer.seek(0)
                            doc.add_picture(img_buffer)
                            
                        # Закрываем буфер
                        img_buffer.close()
                        
                    except Exception as e:
                        logger.error(f"Ошибка сохранения изображения: {e}")
            
            # Добавляем разделитель между вопросами
            if idx < len(self.questions) - 1:
                separator = '~' * 80
                p = doc.add_paragraph(separator)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            self.progress.emit(int((idx + 1) / len(self.questions) * 100))
        
        doc.save(self.export_path)
    
    def export_to_txt(self):
        """Экспорт в TXT с сохранением форматирования."""
        with open(self.export_path, 'w', encoding='utf-8') as f:
            for idx, question in enumerate(self.questions):
                # Получаем текст вопроса
                question_text = question['text']
                
                # Если выбрана нумерация по порядку, заменяем номер в тексте
                if self.export_options.get('numbering') == 'sequential':
                    # Ищем и заменяем "Aufgabe XX" на "Aufgabe [номер]"
                    pattern = r'(Aufgabe\s+)(\d+)'
                    new_number = str(idx + 1)
                    question_text = re.sub(pattern, rf'\g<1>{new_number}', question_text, flags=re.IGNORECASE)
                
                # Сохраняем текст с оригинальным форматированием
                text = question_text.replace('[BILD]', '[Изображение]')
                f.write(text + '\n')
                
                if question['images']:
                    f.write(f"Количество изображений: {len(question['images'])}\n")
                
                f.write('\n' + '='*50 + '\n\n')
                
                self.progress.emit(int((idx + 1) / len(self.questions) * 100))
    
    def export_to_html(self):
        """Экспорт в HTML."""
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Вопросы</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; }
                .question { margin-bottom: 40px; border-bottom: 1px solid #ddd; padding-bottom: 20px; }
                .question-text { margin-bottom: 15px; line-height: 1.6; white-space: pre-wrap; }
                .image-container { margin: 15px 0; text-align: center; }
                .image-container img { max-width: 600px; max-height: 400px; border: 1px solid #ddd; }
                .separator { border-top: 2px dashed #ccc; margin: 30px 0; }
            </style>
        </head>
        <body>
            <h1>Список вопросов</h1>
        """
        
        # Высота страницы A4 в пикселях (при 96 DPI) и 25% от нее
        PAGE_HEIGHT_PIXELS = int(11.69 * 96)  # ~1123 пикселей
        MAX_HEIGHT_PIXELS = int(PAGE_HEIGHT_PIXELS * 0.25)  # ~281 пикселей
        
        for idx, question in enumerate(self.questions):
            html_content += f"""
            <div class="question">
            """
            
            # Получаем текст вопроса
            question_text = question['text']
            
            # Если выбрана нумерация по порядку, заменяем номер в тексте
            if self.export_options.get('numbering') == 'sequential':
                # Ищем и заменяем "Aufgabe XX" на "Aufgabe [номер]"
                pattern = r'(Aufgabe\s+)(\d+)'
                new_number = str(idx + 1)
                question_text = re.sub(pattern, rf'\g<1>{new_number}', question_text, flags=re.IGNORECASE)
            
            html_content += '<div class="question-text">'
            
            text_parts = question_text.split('[BILD]')
            images = question['images']
            
            for i, part in enumerate(text_parts):
                if part.strip():
                    # Сохраняем форматирование с заменой переносов строк
                    part_html = part.replace('\n', '<br>')
                    html_content += part_html
                
                if i < len(images):
                    try:
                        import base64
                        img_bytes = images[i]
                        
                        # Загружаем изображение для проверки размеров
                        img = Image.open(io.BytesIO(img_bytes))
                        img_width, img_height = img.size
                        
                        # Если высота изображения больше 25% страницы, уменьшаем
                        if img_height > MAX_HEIGHT_PIXELS:
                            scale_factor = MAX_HEIGHT_PIXELS / img_height
                            new_width = int(img_width * scale_factor)
                            new_height = MAX_HEIGHT_PIXELS
                            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        
                        # Сохраняем изображение в буфер
                        img_buffer = io.BytesIO()
                        img.save(img_buffer, format='PNG')
                        img_buffer.seek(0)
                        img_bytes = img_buffer.getvalue()
                        img_buffer.close()
                        
                        img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                        html_content += f"""
                        <div class="image-container">
                            <img src="data:image/png;base64,{img_base64}" 
                                 alt="Изображение {i+1}">
                        </div>
                        """
                    except Exception as e:
                        logger.error(f"Ошибка обработки изображения для HTML: {e}")
                        html_content += f'<div>[Изображение {i+1}]</div>'
            
            html_content += """
                </div>
            </div>
            """
            
            if idx < len(self.questions) - 1:
                html_content += '<div class="separator"></div>'
            
            self.progress.emit(int((idx + 1) / len(self.questions) * 100))
        
        html_content += """
        </body>
        </html>
        """
        
        with open(self.export_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def export_to_json(self):
        """Экспорт в JSON."""
        import json
        import base64
        
        data = {
            'total_questions': len(self.questions),
            'questions': []
        }
        
        # Высота страницы A4 в пикселях (при 96 DPI) и 25% от нее
        PAGE_HEIGHT_PIXELS = int(11.69 * 96)  # ~1123 пикселей
        MAX_HEIGHT_PIXELS = int(PAGE_HEIGHT_PIXELS * 0.25)  # ~281 пикселей
        
        for idx, question in enumerate(self.questions):
            # Получаем текст вопроса
            question_text = question['text']
            
            # Если выбрана нумерация по порядку, заменяем номер в тексте
            if self.export_options.get('numbering') == 'sequential':
                # Ищем и заменяем "Aufgabe XX" на "Aufgabe [номер]"
                pattern = r'(Aufgabe\s+)(\d+)'
                new_number = str(idx + 1)
                question_text = re.sub(pattern, rf'\g<1>{new_number}', question_text, flags=re.IGNORECASE)
            
            question_data = {
                'number': idx + 1 if self.export_options.get('numbering') == 'sequential' else None,
                'text': question_text,
                'images_count': len(question['images']),
                'images': []
            }
            
            # Конвертируем изображения в base64
            for img_idx, img_bytes in enumerate(question['images']):
                try:
                    # Загружаем изображение для проверки размеров
                    img = Image.open(io.BytesIO(img_bytes))
                    img_width, img_height = img.size
                    
                    # Если высота изображения больше 25% страницы, уменьшаем
                    if img_height > MAX_HEIGHT_PIXELS:
                        scale_factor = MAX_HEIGHT_PIXELS / img_height
                        new_width = int(img_width * scale_factor)
                        new_height = MAX_HEIGHT_PIXELS
                        img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                    
                    # Сохраняем изображение в буфер
                    img_buffer = io.BytesIO()
                    img.save(img_buffer, format='PNG')
                    img_buffer.seek(0)
                    img_bytes = img_buffer.getvalue()
                    img_buffer.close()
                    
                    img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                    question_data['images'].append({
                        'index': img_idx + 1,
                        'data': img_base64,
                        'size': len(img_bytes)
                    })
                except Exception as e:
                    logger.error(f"Ошибка кодирования изображения: {e}")
            
            data['questions'].append(question_data)
            self.progress.emit(int((idx + 1) / len(self.questions) * 100))
        
        with open(self.export_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    
    def export_to_pdf(self):
        """Экспорт в PDF."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.units import cm
            
            # Регистрируем шрифт
            try:
                # Попробуем найти стандартные шрифты
                font_paths = [
                    'C:/Windows/Fonts/times.ttf',  # Windows
                    '/usr/share/fonts/truetype/times.ttf',  # Linux
                    '/Library/Fonts/Times New Roman.ttf'  # Mac
                ]
                
                font_found = False
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        pdfmetrics.registerFont(TTFont('Times', font_path))
                        font_name = 'Times'
                        font_found = True
                        break
                
                if not font_found:
                    # Используем стандартный шрифт
                    font_name = 'Helvetica'
            except:
                font_name = 'Helvetica'
            
            # Создаем PDF с полями
            c = canvas.Canvas(self.export_path, pagesize=A4)
            width, height = A4
            
            # Устанавливаем поля
            left_margin = 3 * cm
            right_margin = 1 * cm
            top_margin = 1 * cm
            bottom_margin = 1 * cm
            
            # Доступная ширина и высота
            available_width = width - left_margin - right_margin
            available_height = height - top_margin - bottom_margin
            
            y_position = height - top_margin
            page_num = 1
            
            # Высота страницы A4 в пикселях (при 72 DPI) и 25% от нее
            PAGE_HEIGHT_PIXELS = int(11.69 * 72)  # ~842 пикселей
            MAX_HEIGHT_PIXELS = int(PAGE_HEIGHT_PIXELS * 0.25)  # ~210 пикселей
            
            for idx, question in enumerate(self.questions):
                # Проверяем, нужна ли новая страница
                if y_position < bottom_margin + 100:
                    c.showPage()
                    y_position = height - top_margin
                    page_num += 1
                
                # Получаем текст вопроса
                question_text = question['text']
                
                # Если выбрана нумерация по порядку, заменяем номер в тексте
                if self.export_options.get('numbering') == 'sequential':
                    # Ищем и заменяем "Aufgabe XX" на "Aufgabe [номер]"
                    pattern = r'(Aufgabe\s+)(\d+)'
                    new_number = str(idx + 1)
                    question_text = re.sub(pattern, rf'\g<1>{new_number}', question_text, flags=re.IGNORECASE)
                
                # Удаляем маркеры изображений из текста
                text_parts = question_text.split('[BILD]')
                images = question['images']
                
                # Обрабатываем текст и изображения
                for i, text_part in enumerate(text_parts):
                    if text_part.strip():
                        # Обрабатываем текст
                        lines = text_part.split('\n')
                        for line in lines:
                            if line.strip():
                                # Проверяем, нужна ли новая страница
                                if y_position < bottom_margin + 20:
                                    c.showPage()
                                    y_position = height - top_margin
                                    page_num += 1
                                
                                c.setFont(font_name, 11)
                                c.drawString(left_margin, y_position - 15, line)
                                y_position -= 15
                            else:
                                # Пустая строка
                                y_position -= 10
                    
                    # Добавляем изображение
                    if i < len(images):
                        if y_position < bottom_margin + 150:
                            c.showPage()
                            y_position = height - top_margin
                            page_num += 1
                        
                        try:
                            img = Image.open(io.BytesIO(images[i]))
                            img_width, img_height = img.size
                            
                            # Если высота изображения больше 25% страницы, уменьшаем
                            if img_height > MAX_HEIGHT_PIXELS:
                                scale_factor = MAX_HEIGHT_PIXELS / img_height
                                new_width = int(img_width * scale_factor)
                                new_height = MAX_HEIGHT_PIXELS
                                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                                img_width, img_height = new_width, new_height
                            
                            # Максимальные размеры для изображения
                            max_img_width = available_width
                            max_img_height = available_height * 0.5
                            
                            # Вычисляем масштаб
                            scale = min(max_img_width/img_width, max_img_height/img_height, 1.0)
                            display_width = img_width * scale
                            display_height = img_height * scale
                            
                            # Проверяем, помещается ли изображение
                            if y_position - display_height < bottom_margin:
                                c.showPage()
                                y_position = height - top_margin
                                page_num += 1
                            
                            # Сохраняем изображение в буфер
                            img_buffer = io.BytesIO()
                            img.save(img_buffer, format='PNG')
                            img_buffer.seek(0)
                            
                            # Добавляем изображение в PDF
                            c.drawImage(ImageReader(img_buffer), left_margin, y_position - display_height, 
                                      width=display_width, height=display_height)
                            
                            y_position -= display_height + 20
                            img_buffer.close()
                            
                        except Exception as e:
                            logger.error(f"Ошибка добавления изображения в PDF: {e}")
                            c.drawString(left_margin, y_position - 20, f"[Изображение {i + 1}]")
                            y_position -= 30
                
                # Добавляем разделитель между вопросами
                if idx < len(self.questions) - 1:
                    if y_position < bottom_margin + 30:
                        c.showPage()
                        y_position = height - top_margin
                        page_num += 1
                    
                    # Линия разделителя
                    c.setStrokeColorRGB(0.8, 0.8, 0.8)
                    c.setLineWidth(0.5)
                    c.line(left_margin, y_position - 20, left_margin + available_width, y_position - 20)
                    c.setStrokeColorRGB(0, 0, 0)
                    
                    y_position -= 40
                
                self.progress.emit(int((idx + 1) / len(self.questions) * 100))
            
            c.save()
            
        except ImportError:
            # Если reportlab не установлен
            raise Exception("Для экспорта в PDF требуется установить reportlab и Pillow: pip install reportlab Pillow")
        except Exception as e:
            logger.error(f"Ошибка при экспорте в PDF: {str(e)}")
            raise


class DragDropLabel(QLabel):
    """QLabel с поддержкой drag-and-drop и темной темой."""
    fileDropped = pyqtSignal(str)
    
    def __init__(self, text, theme_manager):
        super().__init__(text)
        self.theme_manager = theme_manager
        self.setAcceptDrops(True)
        self.setAlignment(Qt.AlignCenter)
        self.update_style()
    
    def update_style(self):
        """Обновляет стиль в соответствии с текущей темой."""
        theme = self.theme_manager.get_theme()
        self.setStyleSheet(f"""
            QLabel {{
                border: 2px dashed {theme['drag_drop_border']};
                border-radius: 10px;
                background-color: {theme['drag_drop_bg']};
                padding: 20px;
                font-size: 14px;
                color: {theme['drag_drop_text']};
            }}
            QLabel:hover {{
                border-color: {theme['accent']};
                background-color: {theme['tertiary_bg']};
            }}
        """)
    
    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
    
    def dropEvent(self, event: QDropEvent):
        for url in event.mimeData().urls():
            file_path = url.toLocalFile()
            if file_path.lower().endswith('.docx'):
                self.fileDropped.emit(file_path)
                break
        event.acceptProposedAction()


class ClickableQuestionWidget(QWidget):
    """Виджет вопроса, который можно выбрать кликом по любой области."""
    clicked = pyqtSignal(int)  # Сигнал с индексом вопроса
    
    def __init__(self, index, parent=None):
        super().__init__(parent)
        self.index = index
        self.setMouseTracking(True)
        
    def mousePressEvent(self, event):
        """Обработка клика по виджету вопроса."""
        self.clicked.emit(self.index)
        super().mousePressEvent(event)


class QuestionWidget(QWidget):
    """Виджет для отображения одного вопроса с чекбоксом."""
    
    def __init__(self, question_data, index, image_cache=None, parent=None):
        super().__init__(parent)
        self.question_data = question_data
        self.index = index
        self.image_cache = image_cache
        self.setup_ui()
    
    def setup_ui(self):
        # Основной контейнер с возможностью клика
        self.clickable_widget = ClickableQuestionWidget(self.index)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.clickable_widget)
        
        # Layout для контента внутри кликабельного виджета
        content_layout = QVBoxLayout(self.clickable_widget)
        content_layout.setContentsMargins(10, 10, 10, 10)
        
        # Основной горизонтальный layout
        main_layout = QHBoxLayout()
        main_layout.setContentsMargins(0, 0, 0, 0)
        
        # Чекбокс
        self.checkbox = QCheckBox()
        self.checkbox.setFixedSize(30, 30)
        self.checkbox.stateChanged.connect(self.on_checkbox_changed)
        main_layout.addWidget(self.checkbox)
        
        # Контент вопроса
        content_widget = QWidget()
        content_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        content_layout_question = QVBoxLayout(content_widget)
        content_layout_question.setContentsMargins(5, 0, 0, 0)
        
        # Разделяем текст на части по маркерам [BILD]
        text_parts = self.question_data['text'].split('[BILD]')
        images = self.question_data['images']
        
        # Отображаем текст и изображения в правильном порядке
        for i, text_part in enumerate(text_parts):
            if text_part.strip():
                # Создаем QLabel для отображения текста с сохранением форматирования
                text_label = QLabel(text_part)
                text_label.setWordWrap(True)
                text_label.setTextInteractionFlags(Qt.TextSelectableByMouse)
                text_label.setTextFormat(Qt.PlainText)  # Сохраняем оригинальное форматирование
                
                # Делаем QLabel прозрачным для событий мыши
                text_label.setAttribute(Qt.WA_TransparentForMouseEvents, True)
                
                content_layout_question.addWidget(text_label)
            
            # Добавляем изображение, если оно есть для этого маркера
            if i < len(images):
                img_bytes = images[i]
                try:
                    if self.image_cache:
                        # Используем кэш изображений
                        pixmap, _ = self.image_cache.scale_and_cache_image(img_bytes)
                    else:
                        # Старый метод без кэша
                        image = Image.open(io.BytesIO(img_bytes))
                        max_width, max_height = 600, 400
                        width, height = image.size
                        if width > max_width or height > max_height:
                            ratio = min(max_width/width, max_height/height)
                            new_size = (int(width * ratio), int(height * ratio))
                            image = image.resize(new_size, Image.Resampling.LANCZOS)
                        
                        if image.mode == "RGB":
                            qimage = QImage(image.tobytes(), image.size[0], image.size[1], QImage.Format_RGB888)
                        elif image.mode == "RGBA":
                            qimage = QImage(image.tobytes(), image.size[0], image.size[1], QImage.Format_RGBA8888)
                        else:
                            image = image.convert("RGB")
                            qimage = QImage(image.tobytes(), image.size[0], image.size[1], QImage.Format_RGB888)
                        
                        pixmap = QPixmap.fromImage(qimage)
                    
                    if pixmap:
                        img_label = QLabel()
                        img_label.setPixmap(pixmap)
                        img_label.setAlignment(Qt.AlignCenter)
                        img_label.setScaledContents(False)
                        img_label.setMaximumSize(600, 400)
                        img_label.setAttribute(Qt.WA_TransparentForMouseEvents, True)
                        
                        image_container = QWidget()
                        image_container_layout = QVBoxLayout(image_container)
                        image_container_layout.setAlignment(Qt.AlignCenter)
                        image_container_layout.addWidget(img_label)
                        image_container.setAttribute(Qt.WA_TransparentForMouseEvents, True)
                        
                        content_layout_question.addWidget(image_container)
                    
                except Exception as e:
                    logger.error(f"Ошибка загрузки изображения: {str(e)}")
                    error_label = QLabel(f"Не удалось загрузить изображение")
                    error_label.setAttribute(Qt.WA_TransparentForMouseEvents, True)
                    content_layout_question.addWidget(error_label)
        
        # Делаем content_widget прозрачным для событий мыши
        content_widget.setAttribute(Qt.WA_TransparentForMouseEvents, True)
        
        main_layout.addWidget(content_widget, 1)
        content_layout.addLayout(main_layout)
        
        # Разделитель с правильным цветом
        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setFrameShadow(QFrame.Sunken)
        layout.addWidget(separator)
    
    def on_checkbox_changed(self, state):
        """Обновляет стиль при изменении состояния чекбокса."""
        theme = self.get_theme()
        if state == Qt.Checked:
            self.clickable_widget.setStyleSheet(f"""
                ClickableQuestionWidget {{
                    background-color: {theme['selected_bg']};
                    border: 2px solid {theme['accent']};
                    border-radius: 5px;
                }}
                ClickableQuestionWidget:hover {{
                    background-color: {theme['hover_bg']};
                }}
            """)
        else:
            self.clickable_widget.setStyleSheet(f"""
                ClickableQuestionWidget {{
                    border: 1px solid transparent;
                    border-radius: 5px;
                }}
                ClickableQuestionWidget:hover {{
                    background-color: {theme['hover_bg']};
                    border-color: {theme['accent']};
                }}
            """)
    
    def get_theme(self):
        """Получает текущую тему из родительского окна."""
        parent = self.parent()
        while parent:
            if hasattr(parent, 'theme_manager'):
                return parent.theme_manager.get_theme()
            parent = parent.parent()
        return ThemeManager().get_theme('light')
    
    def is_checked(self):
        return self.checkbox.isChecked()
    
    def set_checked(self, checked):
        self.checkbox.setChecked(checked)
        self.on_checkbox_changed(Qt.Checked if checked else Qt.Unchecked)
    
    def connect_click_handler(self, handler):
        """Подключает обработчик клика по виджету."""
        self.clickable_widget.clicked.connect(handler)


class ExportDialog(QDialog):
    """Диалог экспорта вопросов."""
    
    def __init__(self, parent=None, ui_texts=None):
        super().__init__(parent)
        self.ui_texts = ui_texts or {}
        self.setup_ui()
    
    def setup_ui(self):
        layout = QVBoxLayout(self)
        
        # Формат экспорта
        format_layout = QHBoxLayout()
        format_layout.addWidget(QLabel(f"{self.ui_texts.get('format', 'Формат')}:"))
        
        self.format_combo = QComboBox()
        self.format_combo.addItems(['DOCX', 'PDF', 'TXT', 'HTML', 'JSON'])
        format_layout.addWidget(self.format_combo)
        
        layout.addLayout(format_layout)
        
        # Параметры экспорта
        self.options_group = QGroupBox(self.ui_texts.get('parameters', 'Параметры'))
        options_layout = QVBoxLayout()
        
        self.include_images_check = QCheckBox(self.ui_texts.get('include_images', 'Включать изображения'))
        self.include_images_check.setChecked(True)
        options_layout.addWidget(self.include_images_check)
        
        self.options_group.setLayout(options_layout)
        layout.addWidget(self.options_group)
        
        # Кнопки
        buttons = QDialogButtonBox(
            QDialogButtonBox.Ok | QDialogButtonBox.Cancel,
            Qt.Horizontal, self
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
    
    def get_export_options(self):
        """Возвращает выбранные параметры экспорта."""
        return {
            'format': self.format_combo.currentText().lower(),
            'include_images': self.include_images_check.isChecked(),
            'numbering': 'sequential'  # По умолчанию используется нумерация из основного окна
        }


class QuestionApp(QMainWindow):
    """Главное окно приложения."""
    
    def __init__(self):
        super().__init__()
        
        # Инициализация компонентов
        self.settings = QSettings("QuestionApp", "QuestionApp")
        self.theme_manager = ThemeManager()
        self.image_cache = ImageCache()
        
        # Переменные
        self.questions = []
        self.question_widgets = []
        self.current_file_path = None
        self.random_count = 30  # Количество случайных вопросов по умолчанию
        self.numbering_type = 'sequential'  # Тип нумерации: 'original' или 'sequential'
        
        # Загрузка настроек
        self.load_settings()
        
        # Тексты интерфейса
        self.ui_texts = self.load_ui_texts()
        
        self.setup_ui()
        self.apply_theme(self.current_theme)
        self.resize(1200, 900)
    
    def load_settings(self):
        """Загружает настройки приложения."""
        self.current_theme = self.settings.value("theme", "light")
        self.current_language = self.settings.value("language", "ru")
        self.random_count = int(self.settings.value("random_count", 30))
        self.numbering_type = self.settings.value("numbering_type", "sequential")
        self.geometry = self.settings.value("geometry")
        
        if self.geometry:
            self.restoreGeometry(self.geometry)
    
    def save_settings(self):
        """Сохраняет настройки приложения."""
        self.settings.setValue("theme", self.current_theme)
        self.settings.setValue("language", self.current_language)
        self.settings.setValue("random_count", self.random_count)
        self.settings.setValue("numbering_type", self.numbering_type)
        self.settings.setValue("geometry", self.saveGeometry())
    
    def load_ui_texts(self):
        """Загружает тексты интерфейса."""
        texts = {
            'ru': {
                'window_title': 'Выбор вопросов из Word файла',
                'load_file': '📂 Загрузить файл',
                'drag_drop': '📄 Перетащите сюда .docx файл',
                'selected': 'Выбрано:',
                'loaded': 'Загружено:',
                'questions': 'Вопросы',
                'actions': 'Действия',
                'random_select': '🎲 Случайно выбрать',
                'save_selected': '💾 Сохранить выбранные',
                'select_all': '✅ Выбрать все',
                'deselect_all': '❌ Снять выделение',
                'language': 'Язык',
                'theme': 'Тема',
                'export': '📤 Экспорт',
                'search': '🔍 Поиск...',
                'file_dialog_title': 'Выберите Word файл',
                'success_load': 'Успех',
                'error_load': 'Ошибка',
                'success_save': 'Успех',
                'error_save': 'Ошибка',
                'no_questions': 'Нет выбранных вопросов',
                'less_than_count': 'В файле {} вопросов, меньше чем {}',
                'file_saved': 'Сохранено {} вопросов',
                'export_dialog_title': 'Экспорт вопросов',
                'export_progress': 'Экспорт...',
                'export_complete': 'Экспорт завершен',
                'export_error': 'Ошибка экспорта',
                'loading_file': 'Загрузка файла...',
                'parsing_questions': 'Парсинг вопросов...',
                'creating_widgets': 'Создание интерфейса...',
                'cancel': 'Отмена',
                'settings': 'Настройки',
                'statistics': 'Статистика',
                'selection': 'Выбор',
                'format': 'Формат',
                'numbering': 'Нумерация',
                'parameters': 'Параметры',
                'include_images': 'Включать изображения',
                'original_numbering': 'Как в оригинале',
                'sequential_numbering': 'По порядку',
                'random_count_label': 'Количество:',
                'file_menu': 'Файл',
                'edit_menu': 'Правка',
                'settings_menu': 'Настройки',
                'help_menu': 'Помощь',
                'load_action': 'Загрузить файл',
                'save_action': 'Сохранить выбранные',
                'export_action': 'Экспорт...',
                'exit_action': 'Выход',
                'select_all_action': 'Выбрать все',
                'deselect_all_action': 'Снять выделение',
                'random_action': 'Случайные вопросы',
                'theme_menu': 'Тема оформления',
                'light_theme': 'Светлая',
                'dark_theme': 'Темная',
                'about_action': 'О программе',
                'about_title': 'О программе',
                'about_text': 'QuestionApp\n\nПриложение для работы с вопросами из Word файлов.\nВерсия 2.0\n\nВозможности:\n• Загрузка DOCX файлов с вопросами\n• Отображение текста и изображений\n• Выбор вопросов\n• Экспорт в различные форматы\n• Поддержка тем оформления\n\n Für meine Lieblingslehrerin Andrea Retter\n\n © Sienin Oleksandr 2026',
                'original': 'Оригинал',
                'sequential': 'По порядку',
                'export_success': 'Экспорт успешно завершен',
                'export_failed': 'Ошибка экспорта',
                'no_file_selected': 'Файл не выбран',
                'loading': 'Загрузка...',
                'saving': 'Сохранение...',
                'processing': 'Обработка...',
                'complete': 'Завершено',
                'select_file': 'Выберите файл',
                'file_not_found': 'Файл не найден',
                'invalid_format': 'Неверный формат файла',
                'confirm_exit': 'Подтвердите выход',
                'exit_message': 'Вы уверены, что хотите выйти?',
                'yes': 'Да',
                'no': 'Нет',
                'ok': 'OK',
                'close': 'Закрыть',
                'save': 'Сохранить',
                'open': 'Открыть',
                'delete': 'Удалить',
                'edit': 'Редактировать',
                'view': 'Просмотр',
                'help': 'Помощь',
                'about': 'О программе',
                'preferences': 'Настройки',
                'exit': 'Выход',
                'back': 'Назад',
                'next': 'Далее',
                'finish': 'Завершить',
                'browse': 'Обзор',
                'create': 'Создать',
                'remove': 'Удалить',
                'add': 'Добавить',
                'clear': 'Очистить',
                'reset': 'Сбросить',
                'apply': 'Применить',
                'accept': 'Принять',
                'reject': 'Отклонить'
            },
            'uk': {
                'window_title': 'Вибір питань з Word файлу',
                'load_file': '📂 Завантажити файл',
                'drag_drop': '📄 Перетягніть сюди .docх файл',
                'selected': 'Вибрано:',
                'loaded': 'Завантажено:',
                'questions': 'Питання',
                'actions': 'Дії',
                'random_select': '🎲 Вибрати випадково',
                'save_selected': '💾 Зберегти вибрані',
                'select_all': '✅ Вибрати всі',
                'deselect_all': '❌ Зняти виділення',
                'language': 'Мова',
                'theme': 'Тема',
                'export': '📤 Експорт',
                'search': '🔍 Пошук...',
                'file_dialog_title': 'Виберіть Word файл',
                'success_load': 'Успіх',
                'error_load': 'Помилка',
                'success_save': 'Успіх',
                'error_save': 'Помилка',
                'no_questions': 'Немає вибраних питань',
                'less_than_count': 'У файлі {} питань, менше ніж {}',
                'file_saved': 'Збережено {} питань',
                'export_dialog_title': 'Експорт питань',
                'export_progress': 'Експорт...',
                'export_complete': 'Експорт завершено',
                'export_error': 'Помилка експорту',
                'loading_file': 'Завантаження файлу...',
                'parsing_questions': 'Парсинг питань...',
                'creating_widgets': 'Створення інтерфейсу...',
                'cancel': 'Скасувати',
                'settings': 'Налаштування',
                'statistics': 'Статистика',
                'selection': 'Вибір',
                'format': 'Формат',
                'numbering': 'Нумерація',
                'parameters': 'Параметри',
                'include_images': 'Включати зображення',
                'original_numbering': 'Як в оригіналі',
                'sequential_numbering': 'По порядку',
                'random_count_label': 'Кількість:',
                'file_menu': 'Файл',
                'edit_menu': 'Правка',
                'settings_menu': 'Налаштування',
                'help_menu': 'Допомога',
                'load_action': 'Завантажити файл',
                'save_action': 'Зберегти вибрані',
                'export_action': 'Експорт...',
                'exit_action': 'Вихід',
                'select_all_action': 'Вибрати всі',
                'deselect_all_action': 'Зняти виділення',
                'random_action': 'Випадкові питання',
                'theme_menu': 'Тема оформлення',
                'light_theme': 'Світла',
                'dark_theme': 'Темна',
                'about_action': 'Про програму',
                'about_title': 'Про програму',
                'about_text': 'QuestionApp\n\nПрограма для роботи з питаннями з Word файлів.\nВерсія 2.0\n\nМожливості:\n• Завантаження DOCX файлів з питаннями\n• Відображення тексту та зображень\n• Вибір питань\n• Експорт у різні формати\n• Підтримка тем оформлення\n\n Für meine Lieblingslehrerin Andrea Retter\n\n © Sienin Oleksandr 2026',
                'original': 'Оригінал',
                'sequential': 'По порядку',
                'export_success': 'Експорт успішно завершено',
                'export_failed': 'Помилка експорту',
                'no_file_selected': 'Файл не вибрано',
                'loading': 'Завантаження...',
                'saving': 'Збереження...',
                'processing': 'Обробка...',
                'complete': 'Завершено',
                'select_file': 'Виберіть файл',
                'file_not_found': 'Файл не знайдено',
                'invalid_format': 'Невірний формат файлу',
                'confirm_exit': 'Підтвердіть вихід',
                'exit_message': 'Ви впевнені, що хочете вийти?',
                'yes': 'Так',
                'no': 'Ні',
                'ok': 'OK',
                'close': 'Закрити',
                'save': 'Зберегти',
                'open': 'Відкрити',
                'delete': 'Видалити',
                'edit': 'Редагувати',
                'view': 'Перегляд',
                'help': 'Допомога',
                'about': 'Про програму',
                'preferences': 'Налаштування',
                'exit': 'Вихід',
                'back': 'Назад',
                'next': 'Далі',
                'finish': 'Завершити',
                'browse': 'Огляд',
                'create': 'Створити',
                'remove': 'Видалити',
                'add': 'Додати',
                'clear': 'Очистити',
                'reset': 'Скинути',
                'apply': 'Застосувати',
                'accept': 'Прийняти',
                'reject': 'Відхилити'
            },
            'de': {
                'window_title': 'Fragenauswahl aus Word-Datei',
                'load_file': '📂 Datei laden',
                'drag_drop': '📄 .docx-Datei hierher ziehen',
                'selected': 'Ausgewählt:',
                'loaded': 'Geladen:',
                'questions': 'Fragen',
                'actions': 'Aktionen',
                'random_select': '🎲 Zufällig auswählen',
                'save_selected': '💾 Ausgewählte speichern',
                'select_all': '✅ Alle auswählen',
                'deselect_all': '❌ Auswahl aufheben',
                'language': 'Sprache',
                'theme': 'Thema',
                'export': '📤 Export',
                'search': '🔍 Suchen...',
                'file_dialog_title': 'Word-Datei auswählen',
                'success_load': 'Erfolg',
                'error_load': 'Fehler',
                'success_save': 'Erfolg',
                'error_save': 'Fehler',
                'no_questions': 'Keine Fragen ausgewählt',
                'less_than_count': 'Die Datei hat {} Fragen, weniger als {}',
                'file_saved': '{} Fragen gespeichert',
                'export_dialog_title': 'Fragen exportieren',
                'export_progress': 'Exportieren...',
                'export_complete': 'Export abgeschlossen',
                'export_error': 'Exportfehler',
                'loading_file': 'Datei wird geladen...',
                'parsing_questions': 'Fragen werden geparst...',
                'creating_widgets': 'Erstelle Benutzeroberfläche...',
                'cancel': 'Abbrechen',
                'settings': 'Einstellungen',
                'statistics': 'Statistik',
                'selection': 'Auswahl',
                'format': 'Format',
                'numbering': 'Nummerierung',
                'parameters': 'Parameter',
                'include_images': 'Bilder einbeziehen',
                'original_numbering': 'Wie im Original',
                'sequential_numbering': 'In Reihenfolge',
                'random_count_label': 'Anzahl:',
                'file_menu': 'Datei',
                'edit_menu': 'Bearbeiten',
                'settings_menu': 'Einstellungen',
                'help_menu': 'Hilfe',
                'load_action': 'Datei laden',
                'save_action': 'Ausgewählte speichern',
                'export_action': 'Export...',
                'exit_action': 'Beenden',
                'select_all_action': 'Alle auswählen',
                'deselect_all_action': 'Auswahl aufheben',
                'random_action': 'Zufällige Fragen',
                'theme_menu': 'Design',
                'light_theme': 'Hell',
                'dark_theme': 'Dunkel',
                'about_action': 'Über',
                'about_title': 'Über',
                'about_text': 'QuestionApp\n\nAnwendung zur Arbeit mit Fragen aus Word-Dateien.\nVersion 2.0\n\nFunktionen:\n• Laden von DOCX-Dateien mit Fragen\n• Anzeige von Text und Bildern\n• Auswahl von Fragen\n• Export in verschiedene Formate\n• Unterstützung von Designs\n\n Für meine Lieblingslehrerin Andrea Retter\n\n © Sienin Oleksandr 2026',
                'original': 'Original',
                'sequential': 'In Reihenfolge',
                'export_success': 'Export erfolgreich abgeschlossen',
                'export_failed': 'Exportfehler',
                'no_file_selected': 'Keine Datei ausgewählt',
                'loading': 'Laden...',
                'saving': 'Speichern...',
                'processing': 'Verarbeitung...',
                'complete': 'Abgeschlossen',
                'select_file': 'Datei auswählen',
                'file_not_found': 'Datei nicht gefunden',
                'invalid_format': 'Ungültiges Dateiformat',
                'confirm_exit': 'Beenden bestätigen',
                'exit_message': 'Sind Sie sicher, dass Sie beenden möchten?',
                'yes': 'Ja',
                'no': 'Nein',
                'ok': 'OK',
                'close': 'Schließen',
                'save': 'Speichern',
                'open': 'Öffnen',
                'delete': 'Löschen',
                'edit': 'Bearbeiten',
                'view': 'Anzeigen',
                'help': 'Hilfe',
                'about': 'Über',
                'preferences': 'Einstellungen',
                'exit': 'Beenden',
                'back': 'Zurück',
                'next': 'Weiter',
                'finish': 'Fertigstellen',
                'browse': 'Durchsuchen',
                'create': 'Erstellen',
                'remove': 'Entfernen',
                'add': 'Hinzufügen',
                'clear': 'Löschen',
                'reset': 'Zurücksetzen',
                'apply': 'Anwenden',
                'accept': 'Akzeptieren',
                'reject': 'Ablehnen'
            }
        }
        
        return texts.get(self.current_language, texts['ru'])
    
    def retranslate_ui(self):
        """Полностью переводит интерфейс."""
        # Перезагружаем тексты
        self.ui_texts = self.load_ui_texts()
        
        # Обновляем все элементы интерфейса
        self.update_ui_texts()
        
        # Обновляем меню
        if hasattr(self, 'file_menu'):
            self.file_menu.setTitle(self.ui_texts['file_menu'])
            self.load_action.setText(self.ui_texts['load_action'])
            self.save_action.setText(self.ui_texts['save_action'])
            self.export_action.setText(self.ui_texts['export_action'])
            self.exit_action.setText(self.ui_texts['exit_action'])
        
        if hasattr(self, 'edit_menu'):
            self.edit_menu.setTitle(self.ui_texts['edit_menu'])
            self.select_all_action.setText(self.ui_texts['select_all_action'])
            self.deselect_all_action.setText(self.ui_texts['deselect_all_action'])
            self.random_action.setText(self.ui_texts['random_action'])
        
        if hasattr(self, 'settings_menu'):
            self.settings_menu.setTitle(self.ui_texts['settings_menu'])
            self.theme_menu.setTitle(self.ui_texts['theme_menu'])
            self.light_action.setText(self.ui_texts['light_theme'])
            self.dark_action.setText(self.ui_texts['dark_theme'])
        
        if hasattr(self, 'help_menu'):
            self.help_menu.setTitle(self.ui_texts['help_menu'])
            self.about_action.setText(self.ui_texts['about_action'])
        
        # Обновляем комбобокс языков
        if hasattr(self, 'lang_combo'):
            self.lang_combo.blockSignals(True)
            self.lang_combo.clear()
            self.lang_combo.addItems(['🇷🇺 Русский', '🇺🇦 Українська', '🇩🇪 Deutsch'])
            
            # Устанавливаем текущий язык
            lang_map = {'ru': 0, 'uk': 1, 'de': 2}
            if self.current_language in lang_map:
                self.lang_combo.setCurrentIndex(lang_map[self.current_language])
            self.lang_combo.blockSignals(False)
        
        # Обновляем комбобокс тем
        if hasattr(self, 'theme_combo'):
            self.theme_combo.blockSignals(True)
            self.theme_combo.clear()
            self.theme_combo.addItems([self.ui_texts['light_theme'], self.ui_texts['dark_theme']])
            
            # Устанавливаем текущую тему
            theme_map = {'light': 0, 'dark': 1}
            if self.current_theme in theme_map:
                self.theme_combo.setCurrentIndex(theme_map[self.current_theme])
            self.theme_combo.blockSignals(False)
        
        # Обновляем группу нумерации
        if hasattr(self, 'numbering_group'):
            self.numbering_group.setTitle(self.ui_texts['numbering'])
        if hasattr(self, 'numbering_original'):
            self.numbering_original.setText(self.ui_texts['original_numbering'])
        if hasattr(self, 'numbering_sequential'):
            self.numbering_sequential.setText(self.ui_texts['sequential_numbering'])
    
    def setup_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QHBoxLayout(central_widget)
        
        # Левая панель
        left_panel = self.create_left_panel()
        main_layout.addWidget(left_panel)
        
        # Правая панель (список вопросов)
        right_panel = self.create_right_panel()
        main_layout.addWidget(right_panel, 1)
        
        # Создаем меню
        self.create_menu()
        
        # Настройка горячих клавиш
        self.setup_shortcuts()
    
    def create_left_panel(self):
        panel = QWidget()
        panel.setFixedWidth(350)
        
        layout = QVBoxLayout(panel)
        layout.setAlignment(Qt.AlignTop)
        layout.setSpacing(10)
        
        # Панель настроек
        self.settings_group = QGroupBox(self.ui_texts['settings'])
        settings_layout = QVBoxLayout()
        
        # Выбор языка
        lang_layout = QHBoxLayout()
        lang_layout.addWidget(QLabel(f"{self.ui_texts['language']}:"))
        
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(['🇷🇺 Русский', '🇺🇦 Українська', '🇩🇪 Deutsch'])
        
        # Устанавливаем текущий язык
        lang_map = {'ru': 0, 'uk': 1, 'de': 2}
        if self.current_language in lang_map:
            self.lang_combo.setCurrentIndex(lang_map[self.current_language])
        
        self.lang_combo.currentIndexChanged.connect(self.change_language)
        lang_layout.addWidget(self.lang_combo)
        settings_layout.addLayout(lang_layout)
        
        # Выбор темы
        theme_layout = QHBoxLayout()
        theme_layout.addWidget(QLabel(f"{self.ui_texts['theme']}:"))
        
        self.theme_combo = QComboBox()
        self.theme_combo.addItems([self.ui_texts['light_theme'], self.ui_texts['dark_theme']])
        
        # Устанавливаем текущую тему
        theme_map = {'light': 0, 'dark': 1}
        if self.current_theme in theme_map:
            self.theme_combo.setCurrentIndex(theme_map[self.current_theme])
        
        self.theme_combo.currentIndexChanged.connect(self.change_theme)
        theme_layout.addWidget(self.theme_combo)
        settings_layout.addLayout(theme_layout)
        
        self.settings_group.setLayout(settings_layout)
        layout.addWidget(self.settings_group)
        
        # Кнопка загрузки файла
        self.btn_load = QPushButton(self.ui_texts['load_file'])
        self.btn_load.clicked.connect(self.load_file_dialog)
        layout.addWidget(self.btn_load)
        
        # Drag-and-drop область
        self.drop_label = DragDropLabel(self.ui_texts['drag_drop'], self.theme_manager)
        self.drop_label.setFixedHeight(100)
        self.drop_label.fileDropped.connect(self.load_file)
        layout.addWidget(self.drop_label)
        
        layout.addSpacing(20)
        
        # Поиск
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText(self.ui_texts['search'])
        self.search_edit.textChanged.connect(self.filter_questions)
        layout.addWidget(self.search_edit)
        
        # Счетчики
        self.stats_group = QGroupBox(self.ui_texts['statistics'])
        stats_layout = QVBoxLayout()
        
        self.counter_label = QLabel(f"{self.ui_texts['selected']} 0")
        self.counter_label.setAlignment(Qt.AlignCenter)
        stats_layout.addWidget(self.counter_label)
        
        self.loaded_label = QLabel(f"{self.ui_texts['loaded']} 0")
        self.loaded_label.setAlignment(Qt.AlignCenter)
        stats_layout.addWidget(self.loaded_label)
        
        self.stats_group.setLayout(stats_layout)
        layout.addWidget(self.stats_group)
        
        layout.addStretch()
        
        # Настройка нумерации
        self.numbering_group = QGroupBox(self.ui_texts['numbering'])
        numbering_layout = QVBoxLayout()
        
        self.numbering_original = QRadioButton(self.ui_texts['original_numbering'])
        self.numbering_sequential = QRadioButton(self.ui_texts['sequential_numbering'])
        
        # Устанавливаем выбранный тип нумерации
        if self.numbering_type == 'original':
            self.numbering_original.setChecked(True)
        else:
            self.numbering_sequential.setChecked(True)
        
        self.numbering_button_group = QButtonGroup()
        self.numbering_button_group.addButton(self.numbering_original)
        self.numbering_button_group.addButton(self.numbering_sequential)
        
        self.numbering_original.toggled.connect(self.update_numbering_type)
        self.numbering_sequential.toggled.connect(self.update_numbering_type)
        
        numbering_layout.addWidget(self.numbering_original)
        numbering_layout.addWidget(self.numbering_sequential)
        
        self.numbering_group.setLayout(numbering_layout)
        layout.addWidget(self.numbering_group)
        
        # Кнопки выбора
        self.selection_group = QGroupBox(self.ui_texts['selection'])
        selection_layout = QVBoxLayout()
        
        self.btn_select_all = QPushButton(self.ui_texts['select_all'])
        self.btn_select_all.clicked.connect(self.select_all_questions)
        selection_layout.addWidget(self.btn_select_all)
        
        self.btn_deselect_all = QPushButton(self.ui_texts['deselect_all'])
        self.btn_deselect_all.clicked.connect(self.deselect_all_questions)
        selection_layout.addWidget(self.btn_deselect_all)
        
        self.selection_group.setLayout(selection_layout)
        layout.addWidget(self.selection_group)
        
        # Кнопки действий
        self.actions_group = QGroupBox(self.ui_texts['actions'])
        btn_layout = QVBoxLayout()
        
        # Случайный выбор с настройкой количества
        random_layout = QHBoxLayout()
        self.btn_random = QPushButton(self.ui_texts['random_select'])
        self.btn_random.clicked.connect(self.random_select)
        random_layout.addWidget(self.btn_random)
        
        self.random_spin = QSpinBox()
        self.random_spin.setMinimum(1)
        self.random_spin.setMaximum(1000)
        self.random_spin.setValue(self.random_count)
        self.random_spin.valueChanged.connect(self.update_random_count)
        random_layout.addWidget(self.random_spin)
        
        btn_layout.addLayout(random_layout)
        
        self.btn_save = QPushButton(self.ui_texts['save_selected'])
        self.btn_save.clicked.connect(self.save_selected)
        btn_layout.addWidget(self.btn_save)
        
        self.btn_export = QPushButton(self.ui_texts['export'])
        self.btn_export.clicked.connect(self.show_export_dialog)
        btn_layout.addWidget(self.btn_export)
        
        self.actions_group.setLayout(btn_layout)
        layout.addWidget(self.actions_group)
        
        return panel
    
    def create_right_panel(self):
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # Заголовок
        self.title_label = QLabel(self.ui_texts['questions'])
        self.title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.title_label)
        
        # Scroll area для вопросов
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_content = QWidget()
        self.scroll_layout = QVBoxLayout(self.scroll_content)
        self.scroll_layout.setAlignment(Qt.AlignTop)
        self.scroll_area.setWidget(self.scroll_content)
        
        layout.addWidget(self.scroll_area)
        
        return panel
    
    def create_menu(self):
        """Создает главное меню."""
        menubar = self.menuBar()
        
        # Меню Файл
        self.file_menu = menubar.addMenu(self.ui_texts['file_menu'])
        
        self.load_action = QAction(self.ui_texts['load_action'], self)
        self.load_action.triggered.connect(self.load_file_dialog)
        self.load_action.setShortcut('Ctrl+O')
        self.file_menu.addAction(self.load_action)
        
        self.save_action = QAction(self.ui_texts['save_action'], self)
        self.save_action.triggered.connect(self.save_selected)
        self.save_action.setShortcut('Ctrl+S')
        self.file_menu.addAction(self.save_action)
        
        self.export_action = QAction(self.ui_texts['export_action'], self)
        self.export_action.triggered.connect(self.show_export_dialog)
        self.export_action.setShortcut('Ctrl+E')
        self.file_menu.addAction(self.export_action)
        
        self.file_menu.addSeparator()
        
        self.exit_action = QAction(self.ui_texts['exit_action'], self)
        self.exit_action.triggered.connect(self.close)
        self.exit_action.setShortcut('Ctrl+Q')
        self.file_menu.addAction(self.exit_action)
        
        # Меню Правка
        self.edit_menu = menubar.addMenu(self.ui_texts['edit_menu'])
        
        self.select_all_action = QAction(self.ui_texts['select_all_action'], self)
        self.select_all_action.triggered.connect(self.select_all_questions)
        self.select_all_action.setShortcut('Ctrl+A')
        self.edit_menu.addAction(self.select_all_action)
        
        self.deselect_all_action = QAction(self.ui_texts['deselect_all_action'], self)
        self.deselect_all_action.triggered.connect(self.deselect_all_questions)
        self.deselect_all_action.setShortcut('Ctrl+D')
        self.edit_menu.addAction(self.deselect_all_action)
        
        self.random_action = QAction(self.ui_texts['random_action'], self)
        self.random_action.triggered.connect(self.random_select)
        self.random_action.setShortcut('Ctrl+R')
        self.edit_menu.addAction(self.random_action)
        
        # Меню Настройки
        self.settings_menu = menubar.addMenu(self.ui_texts['settings_menu'])
        
        self.theme_menu = self.settings_menu.addMenu(self.ui_texts['theme_menu'])
        
        self.light_action = QAction(self.ui_texts['light_theme'], self)
        self.light_action.triggered.connect(lambda: self.change_theme_by_name('light'))
        self.theme_menu.addAction(self.light_action)
        
        self.dark_action = QAction(self.ui_texts['dark_theme'], self)
        self.dark_action.triggered.connect(lambda: self.change_theme_by_name('dark'))
        self.theme_menu.addAction(self.dark_action)
        
        # Меню Помощь
        self.help_menu = menubar.addMenu(self.ui_texts['help_menu'])
        
        self.about_action = QAction(self.ui_texts['about_action'], self)
        self.about_action.triggered.connect(self.show_about)
        self.help_menu.addAction(self.about_action)
    
    def setup_shortcuts(self):
        """Настраивает горячие клавиши."""
        QShortcut(QKeySequence("Ctrl+F"), self, self.search_edit.setFocus)
        QShortcut(QKeySequence("Esc"), self, self.clear_search)
    
    def apply_theme(self, theme_name):
        """Применяет тему оформления."""
        self.theme_manager.apply_theme(theme_name, QApplication.instance())
        self.current_theme = theme_name
        
        # Обновляем выбранную тему в комбобоксе
        theme_map = {'light': 0, 'dark': 1}
        if theme_name in theme_map:
            self.theme_combo.setCurrentIndex(theme_map[theme_name])
        
        # Обновляем стиль drag-drop label
        if hasattr(self, 'drop_label'):
            self.drop_label.update_style()
    
    def change_theme(self, index):
        """Изменяет тему оформления."""
        themes = ['light', 'dark']
        if 0 <= index < len(themes):
            self.apply_theme(themes[index])
    
    def change_theme_by_name(self, theme_name):
        """Изменяет тему по имени."""
        self.apply_theme(theme_name)
    
    def change_language(self, index):
        """Изменяет язык интерфейса."""
        languages = ['ru', 'uk', 'de']
        if 0 <= index < len(languages):
            self.current_language = languages[index]
            self.retranslate_ui()
            self.save_settings()
    
    def update_numbering_type(self):
        """Обновляет тип нумерации."""
        if self.numbering_original.isChecked():
            self.numbering_type = 'original'
        else:
            self.numbering_type = 'sequential'
        self.save_settings()
    
    def update_random_count(self, value):
        """Обновляет количество случайных вопросов."""
        self.random_count = value
        self.save_settings()
    
    def update_ui_texts(self):
        """Обновляет все тексты интерфейса."""
        if not self.ui_texts:
            return
            
        self.setWindowTitle(self.ui_texts['window_title'])
        
        # Обновляем левую панель
        self.btn_load.setText(self.ui_texts['load_file'])
        self.drop_label.setText(self.ui_texts['drag_drop'])
        self.search_edit.setPlaceholderText(self.ui_texts['search'])
        
        # Обновляем счетчики
        count = sum(1 for widget in self.question_widgets if widget.is_checked())
        self.counter_label.setText(f"{self.ui_texts['selected']}: {count}")
        self.loaded_label.setText(f"{self.ui_texts['loaded']}: {len(self.questions)}")
        
        # Обновляем группы и кнопки
        self.btn_select_all.setText(self.ui_texts['select_all'])
        self.btn_deselect_all.setText(self.ui_texts['deselect_all'])
        self.btn_random.setText(self.ui_texts['random_select'])
        self.btn_save.setText(self.ui_texts['save_selected'])
        self.btn_export.setText(self.ui_texts['export'])
        
        # Обновляем заголовки групп
        self.title_label.setText(self.ui_texts['questions'])
        self.settings_group.setTitle(self.ui_texts['settings'])
        self.stats_group.setTitle(self.ui_texts['statistics'])
        self.selection_group.setTitle(self.ui_texts['selection'])
        self.actions_group.setTitle(self.ui_texts['actions'])
        self.numbering_group.setTitle(self.ui_texts['numbering'])
        
        # Обновляем нумерацию
        self.numbering_original.setText(self.ui_texts['original_numbering'])
        self.numbering_sequential.setText(self.ui_texts['sequential_numbering'])
    
    def load_file_dialog(self):
        """Открывает диалог выбора файла."""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.ui_texts['file_dialog_title'],
            "",
            "Word files (*.docx)"
        )
        
        if file_path:
            self.load_file(file_path)
    
    def load_file(self, file_path):
        """Загружает и парсит файл."""
        try:
            # Показываем прогресс-диалог
            progress = QProgressDialog(
                self.ui_texts['loading_file'],
                self.ui_texts['cancel'],
                0, 100, self
            )
            progress.setWindowModality(Qt.WindowModal)
            progress.setMinimumDuration(0)
            progress.show()
            
            # Загружаем документ
            progress.setLabelText(self.ui_texts['loading_file'])
            doc = Document(file_path)
            progress.setValue(20)
            
            # Парсим вопросы
            progress.setLabelText(self.ui_texts['parsing_questions'])
            self.questions = self.parse_questions_with_images(doc)
            progress.setValue(60)
            
            # Очищаем предыдущие вопросы
            self.clear_questions()
            
            # Создаем виджеты вопросов
            progress.setLabelText(self.ui_texts['creating_widgets'])
            self.create_question_widgets()
            progress.setValue(100)
            
            # Сохраняем путь к файлу
            self.current_file_path = file_path
            
            QMessageBox.information(
                self, 
                self.ui_texts['success_load'], 
                f"Загружено {len(self.questions)} вопросов"
            )
            
        except Exception as e:
            logger.error(f"Ошибка загрузки файла: {str(e)}")
            QMessageBox.critical(
                self, 
                self.ui_texts['error_load'], 
                f"Не удалось загрузить файл:\n{str(e)}"
            )
    
    def parse_questions_with_images(self, doc):
        """Парсит вопросы с изображениями из документа с сохранением форматирования."""
        questions = []
        current_question = {'text': '', 'images': []}
        text_buffer = []
        in_question = False
        
        try:
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text
                
                # Проверяем начало вопроса
                if "---START---" in text:
                    if in_question:  # Предыдущий вопрос не закрыт
                        logger.warning(f"Вопрос не закрыт ---END---")
                        self.save_current_question(current_question, text_buffer, questions)
                    
                    in_question = True
                    current_question = {'text': '', 'images': []}
                    text_buffer = []
                    continue
                
                # Проверяем конец вопроса
                if "---END---" in text:
                    if not in_question:
                        logger.warning(f"Лишний ---END--- в параграфе {i+1}")
                    else:
                        self.save_current_question(current_question, text_buffer, questions)
                    in_question = False
                    continue
                
                if in_question:
                    # Извлекаем изображения из параграфа
                    images = self.extract_images_from_paragraph(paragraph, doc)
                    
                    # Если есть изображения, добавляем их в правильном порядке
                    if images:
                        # Для каждого изображения добавляем текст до него и само изображение
                        # В этом упрощенном подходе мы добавляем весь текст параграфа, затем все изображения
                        # Это может не сохранить точный порядок, если изображения внутри текста
                        if text.strip():
                            text_buffer.append(text)
                        
                        for img_bytes in images:
                            current_question['images'].append(img_bytes)
                            text_buffer.append('[BILD]')
                    else:
                        # Если изображений нет, просто добавляем текст
                        if text:
                            text_buffer.append(text)
            
            # Обрабатываем последний вопрос, если он не закрыт
            if in_question:
                logger.warning(f"Последний вопрос не закрыт ---END---")
                self.save_current_question(current_question, text_buffer, questions)
            
            logger.info(f"Успешно распаршено {len(questions)} вопросов")
            
        except Exception as e:
            logger.error(f"Критическая ошибка при парсинге: {str(e)}")
            raise
        
        return questions
    
    def extract_images_from_paragraph(self, paragraph, doc):
        """Извлекает изображения из параграфа с улучшенной обработкой."""
        images = []
        
        try:
            # Ищем все элементы drawing
            for element in paragraph._element.iter():
                if element.tag.endswith('drawing'):
                    # Ищем blip элементы
                    for blip in element.findall('.//{*}blip'):
                        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rId and rId in doc.part.related_parts:
                            try:
                                image_part = doc.part.related_parts[rId]
                                img_bytes = image_part.blob
                                
                                # Конвертируем изображение в PNG для гарантированной совместимости
                                img = Image.open(io.BytesIO(img_bytes))
                                
                                # Определяем формат
                                if img.format not in ['PNG', 'JPEG', 'GIF', 'BMP']:
                                    # Конвертируем в PNG
                                    png_buffer = io.BytesIO()
                                    if img.mode not in ['RGB', 'RGBA']:
                                        img = img.convert('RGB')
                                    img.save(png_buffer, format='PNG', optimize=True)
                                    img_bytes = png_buffer.getvalue()
                                
                                images.append(img_bytes)
                                
                            except Exception as e:
                                logger.error(f"Ошибка обработки изображения {rId}: {e}")
        
        except Exception as e:
            logger.error(f"Ошибка извлечения изображений: {e}")
        
        return images
    
    def save_current_question(self, current_question, text_buffer, questions):
        """Сохраняет текущий вопрос в список."""
        if text_buffer or current_question['images']:
            question_text = '\n'.join(text_buffer)
            current_question['text'] = question_text
            questions.append(current_question.copy())
    
    def create_question_widgets(self):
        """Создает виджеты вопросов."""
        # Очищаем предыдущие виджеты
        for widget in self.question_widgets:
            widget.deleteLater()
        self.question_widgets.clear()
        
        # Создаем обычные виджеты
        for i, question_data in enumerate(self.questions):
            widget = QuestionWidget(question_data, i, self.image_cache)
            widget.checkbox.stateChanged.connect(self.update_counter)
            widget.connect_click_handler(self.on_question_clicked)
            self.question_widgets.append(widget)
            self.scroll_layout.addWidget(widget)
        
        # Обновляем счетчики
        self.update_loaded_counter()
        self.update_counter()
    
    def clear_questions(self):
        """Очищает список вопросов."""
        for widget in self.question_widgets:
            widget.deleteLater()
        self.question_widgets.clear()
        
        # Очищаем кэш изображений
        self.image_cache.clear_cache()
    
    def on_question_clicked(self, index):
        """Обработчик клика по виджету вопроса."""
        if 0 <= index < len(self.question_widgets):
            widget = self.question_widgets[index]
            widget.set_checked(not widget.is_checked())
    
    def filter_questions(self, text):
        """Фильтрует вопросы по тексту."""
        if not text:
            # Показываем все вопросы
            for widget in self.question_widgets:
                widget.show()
        else:
            # Фильтруем вопросы
            search_text = text.lower()
            for i, widget in enumerate(self.question_widgets):
                if i < len(self.questions):
                    question_text = self.questions[i]['text'].lower()
                    widget.setVisible(search_text in question_text)
    
    def clear_search(self):
        """Очищает поле поиска."""
        self.search_edit.clear()
    
    def select_all_questions(self):
        """Выбирает все вопросы."""
        for widget in self.question_widgets:
            widget.set_checked(True)
        self.update_counter()
    
    def deselect_all_questions(self):
        """Снимает выделение со всех вопросов."""
        for widget in self.question_widgets:
            widget.set_checked(False)
        self.update_counter()
    
    def update_counter(self):
        """Обновляет счетчик выбранных вопросов."""
        count = sum(1 for widget in self.question_widgets if widget.is_checked())
        self.counter_label.setText(f"{self.ui_texts['selected']}: {count}")
    
    def update_loaded_counter(self):
        """Обновляет счетчик загруженных вопросов."""
        self.loaded_label.setText(f"{self.ui_texts['loaded']}: {len(self.questions)}")
    
    def random_select(self):
        """Выбирает случайные вопросы в заданном количестве."""
        count = self.random_spin.value()
        
        if len(self.question_widgets) < count:
            QMessageBox.warning(
                self, 
                self.ui_texts['error_load'], 
                self.ui_texts['less_than_count'].format(len(self.question_widgets), count)
            )
            return
        
        # Сбрасываем все выборы
        for widget in self.question_widgets:
            widget.set_checked(False)
        
        # Выбираем случайные вопросы
        indices = random.sample(range(len(self.question_widgets)), count)
        for idx in indices:
            self.question_widgets[idx].set_checked(True)
        
        self.update_counter()
    
    def save_selected(self):
        """Сохраняет выбранные вопросы в новый DOCX файл."""
        selected_questions = [
            self.questions[i]
            for i, widget in enumerate(self.question_widgets)
            if widget.is_checked()
        ]
        
        if not selected_questions:
            QMessageBox.warning(
                self, 
                self.ui_texts['error_save'], 
                self.ui_texts['no_questions']
            )
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить выбранные вопросы",
            "",
            "Word files (*.docx)"
        )
        
        if not file_path:
            return
        
        # Используем нумерацию из настроек
        self.export_questions(selected_questions, file_path, 'docx', {
            'numbering': self.numbering_type,
            'include_images': True
        })
    
    def show_export_dialog(self):
        """Показывает диалог экспорта."""
        selected_questions = [
            self.questions[i]
            for i, widget in enumerate(self.question_widgets)
            if widget.is_checked()
        ]
        
        if not selected_questions:
            QMessageBox.warning(
                self,
                self.ui_texts['error_save'],
                self.ui_texts['no_questions']
            )
            return
        
        dialog = ExportDialog(self, self.ui_texts)
        if dialog.exec_() == QDialog.Accepted:
            options = dialog.get_export_options()
            
            # Добавляем текущий тип нумерации в опции
            options['numbering'] = self.numbering_type
            
            # Выбираем расширение файла в зависимости от формата
            extensions = {
                'docx': 'Word files (*.docx)',
                'pdf': 'PDF files (*.pdf)',
                'txt': 'Text files (*.txt)',
                'html': 'HTML files (*.html)',
                'json': 'JSON files (*.json)'
            }
            
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                self.ui_texts['export_dialog_title'],
                "",
                extensions.get(options['format'], 'All files (*.*)')
            )
            
            if file_path:
                self.export_questions(selected_questions, file_path, options['format'], options)
    
    def export_questions(self, questions, file_path, export_format, export_options=None):
        """Экспортирует вопросы в указанный формат."""
        if export_options is None:
            export_options = {}
        
        try:
            # Показываем прогресс-диалог
            progress = QProgressDialog(
                self.ui_texts['export_progress'],
                self.ui_texts['cancel'],
                0, 100, self
            )
            progress.setWindowModality(Qt.WindowModal)
            progress.setMinimumDuration(0)
            progress.show()
            
            # Создаем и запускаем поток экспорта
            self.export_worker = ExportWorker(questions, export_format, file_path, export_options)
            self.export_worker.progress.connect(progress.setValue)
            self.export_worker.finished.connect(
                lambda path: self.on_export_finished(path, export_format)
            )
            self.export_worker.error.connect(self.on_export_error)
            
            self.export_worker.start()
            
        except Exception as e:
            logger.error(f"Ошибка начала экспорта: {str(e)}")
            QMessageBox.critical(
                self,
                self.ui_texts['export_error'],
                f"Не удалось начать экспорт:\n{str(e)}"
            )
    
    def on_export_finished(self, file_path, export_format):
        """Обработчик завершения экспорта."""
        QMessageBox.information(
            self,
            self.ui_texts['export_complete'],
            f"Экспорт в {export_format.upper()} завершен:\n{file_path}"
        )
    
    def on_export_error(self, error_message):
        """Обработчик ошибки экспорта."""
        QMessageBox.critical(
            self,
            self.ui_texts['export_error'],
            f"Ошибка экспорта:\n{error_message}"
        )
    
    def show_about(self):
        """Показывает информацию о программе."""
        QMessageBox.about(
            self,
            self.ui_texts['about_title'],
            self.ui_texts['about_text']
        )
    
    def closeEvent(self, event):
        """Обработчик закрытия окна."""
        self.save_settings()
        super().closeEvent(event)


def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    
    window = QuestionApp()
    window.show()
    
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
